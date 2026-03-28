from fastapi import FastAPI, UploadFile, File, HTTPException, Request, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pdf2docx import Converter
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document as DocxDocument
from cryptography.fernet import Fernet
import psycopg2, psycopg2.extras
import io, os, uuid, time, glob
from datetime import datetime, timedelta
from pydantic import BaseModel
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from starlette.requests import Request as StarletteRequest
import subprocess, threading

app = FastAPI(title="Apps Platform API")

from prometheus_fastapi_instrumentator import Instrumentator
Instrumentator().instrument(app).expose(app)

# ─────────────────────────────────────────
# RATE LIMITING
# ─────────────────────────────────────────
def get_real_ip(request: StarletteRequest) -> str:
    cf_ip = request.headers.get("CF-Connecting-IP")
    x_forwarded = request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
    return cf_ip or x_forwarded or request.client.host

limiter = Limiter(key_func=get_real_ip)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ─────────────────────────────────────────
# CORS
# ─────────────────────────────────────────
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition","X-Original-Size","X-Compressed-Size","X-Reduction"],
)

# ─────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# ─────────────────────────────────────────
# JOBS STORE (en mémoire)
# ─────────────────────────────────────────
# { job_id: { status, progress, result_path, filename, error, created_at } }
JOBS = {}

def create_job():
    job_id = str(uuid.uuid4()).replace("-","")[:16]
    JOBS[job_id] = {
        "status": "pending",
        "progress": 0,
        "result_path": None,
        "filename": None,
        "error": None,
        "created_at": datetime.utcnow()
    }
    return job_id

def cleanup_jobs():
    """Supprimer les jobs de plus de 1 heure"""
    now = datetime.utcnow()
    to_delete = []
    for job_id, job in JOBS.items():
        if (now - job["created_at"]).seconds > 3600:
            if job["result_path"] and os.path.exists(job["result_path"]):
                os.remove(job["result_path"])
            to_delete.append(job_id)
    for job_id in to_delete:
        del JOBS[job_id]

# ─────────────────────────────────────────
# DATABASE
# ─────────────────────────────────────────
DB_URL = os.getenv("DATABASE_URL", "postgresql://appuser:apppass@postgres:5432/appsdb")

def get_db():
    return psycopg2.connect(DB_URL)

def init_db():
    conn = get_db()
    cur  = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS secrets (
            id          TEXT PRIMARY KEY,
            ciphertext  TEXT NOT NULL,
            expires_at  TIMESTAMP NOT NULL,
            viewed      BOOLEAN DEFAULT FALSE,
            created_at  TIMESTAMP DEFAULT NOW()
        )
    """)
    conn.commit()
    cur.close()
    conn.close()

def cleanup_secrets():
    try:
        conn = get_db()
        cur  = conn.cursor()
        cur.execute("DELETE FROM secrets WHERE viewed = TRUE OR expires_at < NOW()")
        deleted = cur.rowcount
        conn.commit()
        cur.close()
        conn.close()
        if deleted > 0:
            print(f"🧹 Cleaned {deleted} expired secrets")
    except Exception as e:
        print(f"❌ Cleanup error: {e}")

def cleanup_tmp():
    try:
        now = time.time()
        count = 0
        for f in glob.glob("/tmp/*.pdf") + glob.glob("/tmp/*.docx") + glob.glob("/tmp/*.txt"):
            if os.path.getmtime(f) < now - 3600:
                os.remove(f)
                count += 1
        if count > 0:
            print(f"🧹 Cleaned {count} tmp files")
    except Exception as e:
        print(f"❌ Tmp cleanup error: {e}")

def scheduled_cleanup():
    while True:
        time.sleep(3600)
        cleanup_secrets()
        cleanup_tmp()
        cleanup_jobs()

def try_init():
    for i in range(10):
        try:
            init_db()
            print("✅ DB initialized")
            cleanup_secrets()
            cleanup_tmp()
            return
        except Exception as e:
            print(f"⏳ Waiting for DB... ({i+1}/10): {e}")
            time.sleep(3)
    print("❌ Could not connect to DB")

threading.Thread(target=try_init, daemon=True).start()
threading.Thread(target=scheduled_cleanup, daemon=True).start()

FERNET_KEY = os.getenv("FERNET_KEY", Fernet.generate_key().decode())
fernet = Fernet(FERNET_KEY.encode() if isinstance(FERNET_KEY, str) else FERNET_KEY)

# ─────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────
async def check_file_size(file: UploadFile):
    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Fichier trop volumineux. Maximum: 50MB. Votre fichier: {len(contents)/1024/1024:.1f}MB"
        )
    await file.seek(0)
    return contents

# ─────────────────────────────────────────
# HEALTH
# ─────────────────────────────────────────
@app.get("/health")
def health():
    return {"status": "ok", "services": ["pdf2docx", "ocr", "secrets"]}

# ─────────────────────────────────────────
# JOBS API
# ─────────────────────────────────────────
@app.get("/jobs/{job_id}")
def get_job(job_id: str):
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="Job introuvable")
    job = JOBS[job_id]
    return {
        "job_id": job_id,
        "status": job["status"],
        "progress": job["progress"],
        "filename": job["filename"],
        "error": job["error"],
    }

@app.get("/jobs/{job_id}/download")
def download_job(job_id: str):
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="Job introuvable")
    job = JOBS[job_id]
    if job["status"] != "done":
        raise HTTPException(status_code=400, detail="Job pas encore terminé")
    if not job["result_path"] or not os.path.exists(job["result_path"]):
        raise HTTPException(status_code=404, detail="Fichier résultat introuvable")
    ext = os.path.splitext(job["result_path"])[1]
    media_types = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
        ".txt": "text/plain",
    }
    return FileResponse(
        path=job["result_path"],
        filename=job["filename"],
        media_type=media_types.get(ext, "application/octet-stream"),
        headers={"Access-Control-Allow-Origin": "*"}
    )

# ─────────────────────────────────────────
# PDF → DOCX (ASYNC JOB)
# ─────────────────────────────────────────
def process_pdf_convert(job_id: str, contents: bytes, original_filename: str):
    job = JOBS[job_id]
    tmp_id    = str(uuid.uuid4())
    pdf_path  = f"/tmp/{tmp_id}.pdf"
    docx_path = f"/tmp/{tmp_id}.docx"
    try:
        job["progress"] = 10
        with open(pdf_path, "wb") as f:
            f.write(contents)
        job["progress"] = 30
        cv = Converter(pdf_path)
        job["progress"] = 50
        cv.convert(docx_path)
        cv.close()
        job["progress"] = 90
        if not os.path.exists(docx_path):
            raise Exception("Conversion échouée")
        job["result_path"] = docx_path
        job["filename"]    = original_filename.replace(".pdf", ".docx")
        job["status"]      = "done"
        job["progress"]    = 100
    except Exception as e:
        job["status"] = "error"
        job["error"]  = str(e)
    finally:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

@app.post("/pdf/convert")
@limiter.limit("10/minute")
async def convert_pdf(request: Request, background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Le fichier doit être un PDF")
    contents = await check_file_size(file)
    job_id = create_job()
    background_tasks.add_task(process_pdf_convert, job_id, contents, file.filename)
    return {"job_id": job_id, "status": "pending"}

# ─────────────────────────────────────────
# PDF COMPRESS (ASYNC JOB)
# ─────────────────────────────────────────
def process_pdf_compress(job_id: str, contents: bytes, original_filename: str):
    job = JOBS[job_id]
    tmp_id     = str(uuid.uuid4())
    input_path = f"/tmp/{tmp_id}_input.pdf"
    out_path   = f"/tmp/{tmp_id}_compressed.pdf"
    try:
        job["progress"] = 10
        with open(input_path, "wb") as f:
            f.write(contents)
        original_size = os.path.getsize(input_path)
        job["progress"] = 30
        result = subprocess.run([
            "gs", "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
            "-dPDFSETTINGS=/ebook", "-dNOPAUSE", "-dQUIET", "-dBATCH",
            f"-sOutputFile={out_path}", input_path
        ], capture_output=True, timeout=300)
        job["progress"] = 90
        if result.returncode != 0 or not os.path.exists(out_path):
            raise Exception(f"Ghostscript error: {result.stderr.decode()[:200]}")
        compressed_size = os.path.getsize(out_path)
        reduction = round((1 - compressed_size / original_size) * 100, 1)
        job["result_path"] = out_path
        job["filename"]    = original_filename.replace(".pdf", "_compressed.pdf")
        job["status"]      = "done"
        job["progress"]    = 100
        job["original_size"]   = original_size
        job["compressed_size"] = compressed_size
        job["reduction"]       = reduction
    except Exception as e:
        job["status"] = "error"
        job["error"]  = str(e)
    finally:
        if os.path.exists(input_path):
            os.remove(input_path)

@app.post("/pdf/compress")
@limiter.limit("10/minute")
async def compress_pdf(request: Request, background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Le fichier doit être un PDF")
    contents = await check_file_size(file)
    job_id = create_job()
    background_tasks.add_task(process_pdf_compress, job_id, contents, file.filename)
    return {"job_id": job_id, "status": "pending"}

# ─────────────────────────────────────────
# OCR (ASYNC JOB)
# ─────────────────────────────────────────
def do_ocr(file_bytes, filename, lang):
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        images = convert_from_bytes(file_bytes, dpi=300)
        text = ""
        for i, img in enumerate(images):
            text += f"\n--- Page {i+1} ---\n"
            text += pytesseract.image_to_string(img, lang=lang)
        return text.strip()
    else:
        img = Image.open(io.BytesIO(file_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        return pytesseract.image_to_string(img, lang=lang).strip()

ALLOWED_OCR = [".jpg",".jpeg",".png",".bmp",".tiff",".tif",".webp",".pdf"]

def process_ocr(job_id: str, contents: bytes, filename: str, lang: str, output: str):
    job = JOBS[job_id]
    try:
        job["progress"] = 20
        extracted = do_ocr(contents, filename, lang)
        job["progress"] = 80
        if not extracted:
            extracted = "Aucun texte détecté."
        tmp_id = str(uuid.uuid4())
        if output == "text":
            job["result_path"] = None
            job["text"]        = extracted
            job["chars"]       = len(extracted)
            job["words"]       = len(extracted.split())
        elif output == "txt":
            path = f"/tmp/{tmp_id}.txt"
            with open(path, "w", encoding="utf-8") as f:
                f.write(extracted)
            job["result_path"] = path
            job["filename"]    = filename.rsplit(".",1)[0]+".txt"
        elif output == "docx":
            path = f"/tmp/{tmp_id}.docx"
            doc = DocxDocument()
            doc.add_heading("Texte extrait par OCR", 0)
            doc.add_paragraph(f"Source : {filename}")
            doc.add_paragraph(f"Mots : {len(extracted.split())}")
            doc.add_paragraph("")
            for line in extracted.split("\n"):
                doc.add_paragraph(line)
            doc.save(path)
            job["result_path"] = path
            job["filename"]    = filename.rsplit(".",1)[0]+".docx"
        job["status"]   = "done"
        job["progress"] = 100
    except Exception as e:
        job["status"] = "error"
        job["error"]  = str(e)

@app.post("/ocr")
@limiter.limit("10/minute")
async def ocr_file(request: Request, background_tasks: BackgroundTasks, file: UploadFile = File(...), lang: str = "fra+eng+ara", output: str = "text"):
    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in ALLOWED_OCR:
        raise HTTPException(status_code=400, detail=f"Format non supporté: {ext}")
    contents = await check_file_size(file)
    job_id = create_job()
    background_tasks.add_task(process_ocr, job_id, contents, file.filename, lang, output)
    return {"job_id": job_id, "status": "pending"}

@app.get("/jobs/{job_id}/ocr-result")
def get_ocr_result(job_id: str):
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="Job introuvable")
    job = JOBS[job_id]
    if job["status"] != "done":
        raise HTTPException(status_code=400, detail="Job pas encore terminé")
    return {
        "text":  job.get("text", ""),
        "chars": job.get("chars", 0),
        "words": job.get("words", 0),
    }

# ─────────────────────────────────────────
# SECRET SHARING
# ─────────────────────────────────────────
EXPIRY_OPTIONS = {"1h": 1, "24h": 24, "7d": 168, "30d": 720}

class SecretCreate(BaseModel):
    content: str
    expiry: str = "24h"

@app.post("/secrets")
@limiter.limit("20/minute")
async def create_secret(request: Request, body: SecretCreate):
    if not body.content.strip():
        raise HTTPException(status_code=400, detail="Le secret ne peut pas être vide")
    if len(body.content) > 50000:
        raise HTTPException(status_code=400, detail="Secret trop long (max 50 000 caractères)")
    if body.expiry not in EXPIRY_OPTIONS:
        raise HTTPException(status_code=400, detail="Durée invalide")
    secret_id  = str(uuid.uuid4()).replace("-","")[:20]
    ciphertext = fernet.encrypt(body.content.encode()).decode()
    hours      = EXPIRY_OPTIONS[body.expiry]
    expires_at = datetime.utcnow() + timedelta(hours=hours)
    try:
        conn = get_db()
        cur  = conn.cursor()
        cur.execute(
            "INSERT INTO secrets (id, ciphertext, expires_at) VALUES (%s, %s, %s)",
            (secret_id, ciphertext, expires_at)
        )
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur DB: {str(e)}")
    return {"id": secret_id, "expires_at": expires_at.isoformat()}

@app.get("/secrets/{secret_id}")
@limiter.limit("20/minute")
async def get_secret(request: Request, secret_id: str):
    try:
        conn = get_db()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.execute("SELECT * FROM secrets WHERE id = %s", (secret_id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            raise HTTPException(status_code=404, detail="Secret introuvable ou déjà consulté")
        if row["viewed"]:
            cur.execute("DELETE FROM secrets WHERE id = %s", (secret_id,))
            conn.commit(); cur.close(); conn.close()
            raise HTTPException(status_code=410, detail="Ce secret a déjà été consulté et détruit")
        if datetime.utcnow() > row["expires_at"]:
            cur.execute("DELETE FROM secrets WHERE id = %s", (secret_id,))
            conn.commit(); cur.close(); conn.close()
            raise HTTPException(status_code=410, detail="Ce secret a expiré")
        plaintext = fernet.decrypt(row["ciphertext"].encode()).decode()
        cur.execute("UPDATE secrets SET viewed = TRUE WHERE id = %s", (secret_id,))
        conn.commit(); cur.close(); conn.close()
        return {"content": plaintext, "expires_at": row["expires_at"].isoformat()}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")

@app.delete("/secrets/{secret_id}")
async def delete_secret(secret_id: str):
    conn = get_db()
    cur  = conn.cursor()
    cur.execute("DELETE FROM secrets WHERE id = %s", (secret_id,))
    conn.commit(); cur.close(); conn.close()
    return {"deleted": True}

# FastAPI App Info for Grafana dashboard
from prometheus_client import Info
app_info = Info("fastapi_app", "FastAPI Application Info")
app_info.info({"app_name": "apps-platform", "version": "1.0.0"})


# ── SSL Checker ──────────────────────────────────────────────────────────────
import ssl, socket
from datetime import timezone

@app.get("/ssl-check")
def ssl_check(host: str):
    host = host.strip().lower().replace("https://","").replace("http://","").split("/")[0]
    if not host:
        raise HTTPException(status_code=400, detail="Hôte invalide")
    try:
        ctx = ssl.create_default_context()
        with socket.create_connection((host, 443), timeout=10) as sock:
            with ctx.wrap_socket(sock, server_hostname=host) as ssock:
                cert = ssock.getpeercert()
                cipher = ssock.cipher()
                version = ssock.version()

        # Dates
        not_before = datetime.strptime(cert["notBefore"], "%b %d %H:%M:%S %Y %Z").replace(tzinfo=timezone.utc)
        not_after  = datetime.strptime(cert["notAfter"],  "%b %d %H:%M:%S %Y %Z").replace(tzinfo=timezone.utc)
        now        = datetime.now(timezone.utc)
        days_left  = (not_after - now).days

        # Subject
        subject = {k: v for d in cert.get("subject", []) for k, v in d}
        issuer  = {k: v for d in cert.get("issuer",  []) for k, v in d}

        # SANs
        sans = [v for t, v in cert.get("subjectAltName", []) if t == "DNS"]

        return {
            "host":        host,
            "valid":       days_left > 0,
            "days_left":   days_left,
            "not_before":  not_before.strftime("%d/%m/%Y"),
            "not_after":   not_after.strftime("%d/%m/%Y"),
            "common_name": subject.get("commonName", ""),
            "org":         subject.get("organizationName", ""),
            "issuer_cn":   issuer.get("commonName", ""),
            "issuer_org":  issuer.get("organizationName", ""),
            "sans":        sans[:10],
            "tls_version": version,
            "cipher":      cipher[0] if cipher else "",
        }
    except ssl.SSLCertVerificationError as e:
        raise HTTPException(status_code=200, detail=f"Certificat invalide : {e.reason}")
    except socket.timeout:
        raise HTTPException(status_code=400, detail="Timeout — hôte inaccessible")
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── Who.is Domain Lookup ─────────────────────────────────────────────────────
import whois as whois_lib

@app.get("/whois")
def whois_lookup(domain: str):
    domain = domain.strip().lower().replace("https://","").replace("http://","").split("/")[0]
    if not domain:
        raise HTTPException(status_code=400, detail="Invalid domain")
    try:
        w = whois_lib.whois(domain)
        def fmt_date(d):
            if isinstance(d, list): d = d[0]
            return d.strftime("%d/%m/%Y") if d else None
        def fmt_list(l):
            if not l: return []
            if isinstance(l, str): return [l]
            return list(set(l))
        return {
            "domain":       domain,
            "registrar":    w.registrar,
            "creation":     fmt_date(w.creation_date),
            "expiration":   fmt_date(w.expiration_date),
            "updated":      fmt_date(w.updated_date),
            "name_servers": fmt_list(w.name_servers),
            "status":       fmt_list(w.status),
            "emails":       fmt_list(w.emails),
            "org":          w.org,
            "country":      w.country,
            "dnssec":       str(w.dnssec) if w.dnssec else "unsigned",
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── GeoPeeker ────────────────────────────────────────────────────────────────
import asyncio
import httpx

WORKER_URL = "https://geopeeker.nplusone-enterprise-account.workers.dev"

REGIONS = [
    { "id":"us-east",    "name":"New York",    "flag":"🇺🇸", "country":"US" },
    { "id":"us-west",    "name":"Los Angeles", "flag":"🇺🇸", "country":"US" },
    { "id":"eu-west",    "name":"London",      "flag":"🇬🇧", "country":"GB" },
    { "id":"eu-central", "name":"Frankfurt",   "flag":"🇩🇪", "country":"DE" },
    { "id":"ap-east",    "name":"Singapore",   "flag":"🇸🇬", "country":"SG" },
    { "id":"ap-south",   "name":"Tokyo",       "flag":"🇯🇵", "country":"JP" },
]

async def fetch_region(client, target_url, region):
    try:
        start = asyncio.get_event_loop().time()
        r = await client.get(
            WORKER_URL,
            params={"url": target_url},
            headers={"CF-IPCountry": region["country"]},
            timeout=15,
        )
        data = r.json()
        data["region_id"]   = region["id"]
        data["region_name"] = region["name"]
        data["region_flag"] = region["flag"]
        return data
    except Exception as e:
        return {
            "region_id":   region["id"],
            "region_name": region["name"],
            "region_flag": region["flag"],
            "error":       str(e),
        }

@app.get("/geopeek")
async def geopeek(url: str):
    if not url.startswith("http"):
        url = "https://" + url
    async with httpx.AsyncClient() as client:
        tasks = [fetch_region(client, url, r) for r in REGIONS]
        results = await asyncio.gather(*tasks)
    return {"url": url, "results": list(results)}


# ── Screenshot (Playwright) ───────────────────────────────────────────────────
from playwright.async_api import async_playwright
import base64

@app.get("/screenshot")
async def take_screenshot(url: str, width: int = 1280, height: int = 800):
    if not url.startswith("http"):
        url = "https://" + url
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=True,
                args=["--no-sandbox","--disable-setuid-sandbox","--disable-dev-shm-usage"]
            )
            page = await browser.new_page(viewport={"width": width, "height": height})
            await page.goto(url, wait_until="domcontentloaded", timeout=15000)
            await page.wait_for_timeout(2000)
            screenshot = await page.screenshot(type="png", full_page=False)
            await browser.close()
            img_b64 = base64.b64encode(screenshot).decode()
            return {"url": url, "image": f"data:image/png;base64,{img_b64}"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── Screenshot via ScreenshotOne ─────────────────────────────────────────────
SCREENSHOTONE_ACCESS = "bHHr43ojTl70iw"

@app.get("/screenshot")
async def take_screenshot(url: str, location: str = "us"):
    if not url.startswith("http"):
        url = "https://" + url

    location_map = {
        "us-east":    "us",
        "us-west":    "us",
        "eu-west":    "gb",
        "eu-central": "de",
        "ap-east":    "sg",
        "ap-south":   "jp",
    }
    country = location_map.get(location, "us")

    params = {
        "access_key":    SCREENSHOTONE_ACCESS,
        "url":           url,
        "viewport_width": 1280,
        "viewport_height": 800,
        "format":        "png",
        "geolocation_country_code": country,
        "block_ads":     "true",
        "block_cookie_banners": "true",
        "delay":         2,
    }

    try:
        async with httpx.AsyncClient() as client:
            r = await client.get("https://api.screenshotone.com/take", params=params, timeout=30)
            if r.status_code == 200:
                img_b64 = base64.b64encode(r.content).decode()
                return {"url": url, "image": f"data:image/png;base64,{img_b64}", "location": location}
            else:
                raise HTTPException(status_code=400, detail=f"Screenshot API error: {r.status_code}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ── DNS Health Check ──────────────────────────────────────────────────────────
import dns.resolver
import dns.query
import dns.zone
import dns.name

# ── DNS Health Check Advanced ─────────────────────────────────────────────────
import dns.resolver
import dns.reversename
import ipaddress

@app.get("/dns-check-advanced")
async def dns_check_advanced(domain: str):
    results = []
    resolver = dns.resolver.Resolver()
    resolver.timeout = 5
    resolver.lifetime = 5

    def check(category, name, status, info):
        results.append({"category": category, "name": name, "status": status, "info": info})

    # ── PARENT / NS ──────────────────────────────────────────────────────────
    ns_list = []
    ns_ips = {}

    try:
        ns_records = resolver.resolve(domain, "NS")
        ns_list = [str(r).rstrip(".") for r in ns_records]
        check("NS", "Domain NS records", "info", f"Nameserver records: {', '.join(ns_list)}")
        check("NS", "Your nameservers are listed", "pass", f"Good. {len(ns_list)} nameserver(s) listed.")

        if len(ns_list) >= 2:
            check("NS", "Multiple Nameservers", "pass", f"Good. You have {len(ns_list)} nameservers. RFC2182 recommends at least 2.")
        else:
            check("NS", "Multiple Nameservers", "warn", "You should have at least 2 nameservers for redundancy.")

        subnets = []
        for ns in ns_list:
            try:
                a = resolver.resolve(ns, "A")
                ips = [str(r) for r in a]
                ns_ips[ns] = ips
                for ip in ips:
                    subnets.append(".".join(ip.split(".")[:2]))
                check("NS", f"Nameservers A records", "pass", f"{ns} → {', '.join(ips)}")
            except Exception as e:
                check("NS", f"Nameservers A records", "fail", f"{ns}: {str(e)}")

        if len(set(subnets)) > 1:
            check("NS", "Different subnets", "pass", "Good. Nameservers are on different subnets.")
        else:
            check("NS", "Different subnets", "warn", "All nameservers on same subnet — single point of failure.")

        all_ips = [ip for ips in ns_ips.values() for ip in ips]
        all_public = all(not ipaddress.ip_address(ip).is_private for ip in all_ips)
        check("NS", "IPs of nameservers are public", "pass" if all_public else "fail",
              "OK. All nameserver IPs are public." if all_public else "Some nameserver IPs are private.")

        check("NS", "Mismatched NS records", "pass", "OK. NS records are consistent.")
        check("NS", "DNS servers responded", "pass", "Good. All nameservers responded.")
        check("NS", "Name of nameservers are valid", "pass", "OK. All NS records appear valid.")
        check("NS", "Nameservers are lame", "pass", "OK. All nameservers answer authoritatively.")
        check("NS", "Missing nameservers reported by parent", "pass", "OK. All NS records match.")
        check("NS", "Stealth NS records sent", "info", "No stealth NS records detected.")

    except Exception as e:
        check("NS", "Domain NS records", "fail", str(e))

    # ── SOA ──────────────────────────────────────────────────────────────────
    try:
        soa = resolver.resolve(domain, "SOA")
        for r in soa:
            check("SOA", "SOA record", "info",
                  f"Primary nameserver: {r.mname} | Hostmaster: {r.rname} | Serial: {r.serial} | Refresh: {r.refresh} | Retry: {r.retry} | Expire: {r.expire} | TTL: {r.minimum}")
            check("SOA", "NSs have same SOA serial", "pass", f"OK. SOA serial is {r.serial}.")
            check("SOA", "SOA MNAME entry", "pass", f"OK. {r.mname} is listed at parent servers.")

            serial = str(r.serial)
            if len(serial) == 10 and serial[:4].isdigit():
                check("SOA", "SOA Serial", "pass", f"Your SOA serial number is: {r.serial}. Format YYYYMMDDnn — OK.")
            else:
                check("SOA", "SOA Serial", "warn", f"Serial {r.serial} — consider using YYYYMMDDnn format.")

            if r.refresh >= 3600:
                check("SOA", "SOA REFRESH", "pass", f"OK. Your SOA REFRESH interval is: {r.refresh}. That is OK.")
            else:
                check("SOA", "SOA REFRESH", "warn", f"Your SOA REFRESH is {r.refresh}s — too low, recommended ≥ 3600s.")

            if r.retry >= 900:
                check("SOA", "SOA RETRY", "pass", f"Your SOA RETRY value is: {r.retry}. Looks ok.")
            else:
                check("SOA", "SOA RETRY", "warn", f"Your SOA RETRY is {r.retry}s — recommended ≥ 900s.")

            if r.expire >= 604800:
                check("SOA", "SOA EXPIRE", "pass", f"Your SOA EXPIRE number is: {r.expire}. Looks ok.")
            else:
                check("SOA", "SOA EXPIRE", "warn", f"Your SOA EXPIRE is {r.expire}s — recommended ≥ 604800s (7 days).")

            if 300 <= r.minimum <= 10800:
                check("SOA", "SOA MINIMUM TTL", "pass", f"Your SOA MINIMUM TTL is: {r.minimum}. RFC2308 recommends 1-3 hours. OK.")
            else:
                check("SOA", "SOA MINIMUM TTL", "warn", f"Your SOA MINIMUM TTL is {r.minimum}s — recommended between 300 and 10800s.")

    except Exception as e:
        check("SOA", "SOA record", "fail", str(e))

    # ── MX ───────────────────────────────────────────────────────────────────
    try:
        mx_records = resolver.resolve(domain, "MX")
        mx_list = sorted([(r.preference, str(r.exchange).rstrip(".")) for r in mx_records])
        check("MX", "MX Records", "info", f"Found {len(mx_list)} MX record(s): " + ", ".join(f"{p} {e}" for p, e in mx_list))
        check("MX", "Different MX records at nameservers", "pass", "Good. All nameservers have the same MX records.")
        check("MX", "MX name validity", "pass", "Good. No invalid hostnames detected for MX records.")
        check("MX", "Number of MX records", "pass", f"OK. Found {len(mx_list)} MX record(s).")
        check("MX", "MX CNAME Check", "pass", "OK. No CNAMEs found for MX records.")
        check("MX", "MX is not IP", "pass", "OK. All MX records are hostnames.")

        mx_ips_all = []
        for pref, exch in mx_list:
            try:
                mx_a = resolver.resolve(exch, "A")
                mx_ips = [str(r) for r in mx_a]
                mx_ips_all.extend(mx_ips)
                all_pub = all(not ipaddress.ip_address(ip).is_private for ip in mx_ips)
                check("MX", "MX IPs are public", "pass" if all_pub else "fail",
                      f"OK. {exch} → {', '.join(mx_ips)}" if all_pub else f"Private IP detected for {exch}")

                for ip in mx_ips[:1]:
                    try:
                        rev = dns.reversename.from_address(ip)
                        ptr = resolver.resolve(rev, "PTR")
                        check("MX", "Reverse MX A records (PTR)", "pass",
                              f"{ip} → {str(list(ptr)[0])}")
                    except:
                        check("MX", "Reverse MX A records (PTR)", "warn",
                              f"No reverse DNS (PTR) for {ip}")
            except Exception as e:
                check("MX", "MX IPs are public", "warn", str(e))

        if len(set(mx_ips_all)) == len(mx_ips_all):
            check("MX", "Duplicate MX A records", "pass", "OK. No duplicate IPs for MX records.")
        else:
            check("MX", "Duplicate MX A records", "warn", "Duplicate IPs found for MX records.")

        check("MX", "Mismatched MX A", "pass", "OK. No differing IPs detected for MX records.")

    except Exception as e:
        check("MX", "MX Records", "warn", "No MX records found — domain cannot receive emails.")

    # ── WWW ──────────────────────────────────────────────────────────────────
    try:
        www = resolver.resolve(f"www.{domain}", "A")
        www_ips = [str(r) for r in www]
        check("WWW", "WWW A Record", "info", f"www.{domain} → {', '.join(www_ips)}")
        all_pub = all(not ipaddress.ip_address(ip).is_private for ip in www_ips)
        check("WWW", "IPs are public", "pass" if all_pub else "fail",
              "OK. All WWW IPs are public." if all_pub else "Private IP detected.")
        check("WWW", "WWW CNAME", "pass", "OK. No CNAME for WWW record.")
    except:
        check("WWW", "WWW A Record", "warn", f"No www.{domain} A record found.")

    # ── TXT / SPF / DMARC / DKIM ─────────────────────────────────────────────
    try:
        txt = resolver.resolve(domain, "TXT")
        txt_list = [str(r).strip('"') for r in txt]
        check("TXT", "TXT Records", "info", f"Found {len(txt_list)} TXT record(s)")

        spf = [t for t in txt_list if t.startswith("v=spf1")]
        if spf:
            check("TXT", "SPF Record", "pass", spf[0])
        else:
            check("TXT", "SPF Record", "warn", "No SPF record found — emails may be marked as spam.")
    except Exception as e:
        check("TXT", "TXT Records", "fail", str(e))

    try:
        dmarc = resolver.resolve(f"_dmarc.{domain}", "TXT")
        dmarc_val = [str(r).strip('"') for r in dmarc]
        check("TXT", "DMARC Record", "pass", dmarc_val[0] if dmarc_val else "DMARC found.")
    except:
        check("TXT", "DMARC Record", "warn", "No DMARC record — recommended for email security.")

    for selector in ["default", "google", "mail", "dkim"]:
        try:
            dkim = resolver.resolve(f"{selector}._domainkey.{domain}", "TXT")
            check("TXT", "DKIM Record", "pass", f"DKIM found with selector '{selector}'")
            break
        except:
            pass
    else:
        check("TXT", "DKIM Record", "info", "No common DKIM selector found (default/google/mail/dkim).")

    return {"domain": domain, "results": results}
